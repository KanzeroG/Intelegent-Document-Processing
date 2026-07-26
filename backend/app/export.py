"""Export approved documents to JSON / CSV / XLSX / DOCX / PDF, plus a mock
downstream API.

Deliverable #4: once a human approves an extraction, the clean structured data
is exported to CSV/JSON and handed to a mock API endpoint (standing in for a
real accounting/ERP system). These helpers operate on the stored record dict
(see db.py), whose `data` already matches the ground-truth field names.

Two families of format live here, and they answer different questions:

  * **CSV / JSON** — machine-readable. Columns mirror `Source/ground_truth.csv`
    exactly (line items as an embedded JSON blob) so an export can be diffed
    against the labels or piped into another system.
  * **XLSX / DOCX / PDF** — human-readable. A narrower, formatted column set
    plus a totals row, for the Archive's "give me this month as a report"
    button. Rupiah is rendered with Indonesian dot-thousands grouping.
"""

from __future__ import annotations

import csv
import io
import json
from datetime import date
from typing import Any, Callable

# Column order mirrors Source/ground_truth.csv for a drop-in comparison.
_CSV_COLUMNS = [
    "doc_id", "doc_type", "doc_number", "vendor", "buyer", "doc_date",
    "currency", "line_item_count", "subtotal", "tax_amount", "total_amount",
    "line_items",
]


def _flat_row(rec: dict[str, Any]) -> dict[str, Any]:
    """Flatten a stored record into ground_truth-shaped columns."""
    data = rec.get("data") or {}
    line_items = data.get("line_items", [])
    return {
        "doc_id": rec.get("id"),
        "doc_type": rec.get("doc_type"),
        "doc_number": data.get("doc_number"),
        "vendor": data.get("vendor"),
        "buyer": data.get("buyer"),
        "doc_date": data.get("doc_date"),
        "currency": data.get("currency"),
        "line_item_count": data.get("line_item_count", len(line_items)),
        "subtotal": data.get("subtotal"),
        "tax_amount": data.get("tax_amount"),
        "total_amount": data.get("total_amount"),
        "line_items": json.dumps(line_items, ensure_ascii=False),
    }


def to_json(rec: dict[str, Any]) -> str:
    """Pretty JSON export of one approved record."""
    payload = {
        "doc_id": rec.get("id"),
        "doc_type": rec.get("doc_type"),
        "status": rec.get("status"),
        "confidence": rec.get("confidence"),
        **(rec.get("data") or {}),
    }
    return json.dumps(payload, indent=2, ensure_ascii=False)


def to_csv(rec: dict[str, Any]) -> str:
    """Single-row CSV export (header + row) using ground-truth columns."""
    return to_csv_many([rec])


def to_csv_many(records: list[dict[str, Any]]) -> str:
    """Multi-row CSV export (header + one row per record), ground-truth columns."""
    buf = io.StringIO()
    writer = csv.DictWriter(buf, fieldnames=_CSV_COLUMNS)
    writer.writeheader()
    for rec in records:
        writer.writerow(_flat_row(rec))
    return buf.getvalue()


def to_json_many(records: list[dict[str, Any]]) -> str:
    """Pretty JSON array of records — the multi-doc sibling of `to_json`."""
    return json.dumps(
        [
            {
                "doc_id": rec.get("id"),
                "doc_type": rec.get("doc_type"),
                "status": rec.get("status"),
                "confidence": rec.get("confidence"),
                "approved_at": rec.get("approved_at"),
                **(rec.get("data") or {}),
            }
            for rec in records
        ],
        indent=2,
        ensure_ascii=False,
    )


# --- report formats (XLSX / DOCX / PDF) --------------------------------------

DOC_TYPE_LABEL = {
    "invoice": "Invoice",
    "purchase_order": "Purchase Order",
    "receipt": "Receipt",
}

# Human-readable columns; `money` marks the ones that get Rupiah formatting and
# feed the totals row. Deliberately narrower than _CSV_COLUMNS: a printed report
# has no room for a JSON blob of line items.
_REPORT_COLUMNS: list[tuple[str, str, bool]] = [
    # (header, record key produced by _report_row, is_money)
    ("Approved", "approved_on", False),
    ("Doc ID", "doc_number", False),
    ("Type", "doc_type", False),
    ("Vendor", "vendor", False),
    ("Document date", "doc_date", False),
    ("Subtotal", "subtotal", True),
    ("Tax", "tax_amount", True),
    ("Total", "total_amount", True),
]


def _fmt_amount(value: Any) -> str:
    """Rupiah with Indonesian dot-thousands: 12450000 -> '12.450.000'."""
    if value is None:
        return "—"
    try:
        number = float(value)
    except (TypeError, ValueError):
        return str(value)
    whole = f"{abs(number):,.0f}".replace(",", ".")
    return f"-{whole}" if number < 0 else whole


def _report_row(rec: dict[str, Any]) -> dict[str, Any]:
    """One record reduced to the columns a printed report shows."""
    data = rec.get("data") or {}
    approved = (rec.get("approved_at") or "")[:10]
    return {
        "approved_on": approved or "—",
        # Falls back to the internal id when the model never found a doc number,
        # so every row is still traceable back to a stored document.
        "doc_number": data.get("doc_number") or rec.get("id") or "—",
        "doc_type": DOC_TYPE_LABEL.get(rec.get("doc_type") or "", rec.get("doc_type") or "—"),
        "vendor": data.get("vendor") or "—",
        "doc_date": data.get("doc_date") or "—",
        "subtotal": data.get("subtotal"),
        "tax_amount": data.get("tax_amount"),
        "total_amount": data.get("total_amount"),
        "currency": data.get("currency") or "IDR",
    }


def _money_totals(rows: list[dict[str, Any]]) -> dict[str, float]:
    """Column sums for the money columns; missing values count as zero."""
    totals: dict[str, float] = {}
    for _, key, is_money in _REPORT_COLUMNS:
        if not is_money:
            continue
        totals[key] = sum(float(r[key]) for r in rows if r.get(key) is not None)
    return totals


def to_xlsx(records: list[dict[str, Any]], title: str) -> bytes:
    """Excel workbook: a formatted 'Documents' sheet plus a 'Line Items' sheet.

    Numbers are written as real numbers with an Indonesian display format, so
    the recipient can sum/pivot them rather than re-typing text."""
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font, PatternFill
    from openpyxl.utils import get_column_letter

    rows = [_report_row(r) for r in records]
    wb = Workbook()

    ws = wb.active
    ws.title = "Documents"
    ws.append([title])
    ws["A1"].font = Font(bold=True, size=14)
    ws.append([f"{len(rows)} approved document{'' if len(rows) == 1 else 's'}"])
    ws.append([])

    header_at = ws.max_row + 1
    ws.append([h for h, _, _ in _REPORT_COLUMNS])
    fill = PatternFill("solid", fgColor="E8EAF6")
    for cell in ws[header_at]:
        cell.font = Font(bold=True)
        cell.fill = fill

    # Excel stores format codes in US notation and renders them in the reader's
    # locale — an Indonesian Excel shows "#,##0" as dot-thousands by itself.
    money_format = "#,##0"
    for row in rows:
        ws.append([row[key] for _, key, _ in _REPORT_COLUMNS])
        for col, (_, _, is_money) in enumerate(_REPORT_COLUMNS, start=1):
            if is_money:
                cell = ws.cell(row=ws.max_row, column=col)
                cell.number_format = money_format
                cell.alignment = Alignment(horizontal="right")

    if rows:
        totals = _money_totals(rows)
        total_at = ws.max_row + 2  # one blank row between the data and the total
        for col, (_, key, is_money) in enumerate(_REPORT_COLUMNS, start=1):
            value = "TOTAL" if col == 1 else (totals[key] if is_money else None)
            cell = ws.cell(row=total_at, column=col, value=value)
            cell.font = Font(bold=True)
            if is_money:
                cell.number_format = money_format
                cell.alignment = Alignment(horizontal="right")

    widths = [12, 18, 16, 30, 14, 16, 14, 16]
    for i, width in enumerate(widths, start=1):
        ws.column_dimensions[get_column_letter(i)].width = width
    ws.freeze_panes = ws.cell(row=header_at + 1, column=1)

    # Second sheet: the line items, one row each, keyed back to their document.
    items = wb.create_sheet("Line Items")
    items.append(["Doc ID", "Description", "Qty", "Unit price", "Line total"])
    for cell in items[1]:
        cell.font = Font(bold=True)
        cell.fill = fill
    for rec in records:
        data = rec.get("data") or {}
        label = data.get("doc_number") or rec.get("id")
        for item in data.get("line_items") or []:
            items.append([
                label,
                item.get("description"),
                item.get("qty"),
                item.get("unit_price"),
                item.get("line_total"),
            ])
            for col in (4, 5):
                items.cell(row=items.max_row, column=col).number_format = money_format
    for i, width in enumerate([18, 46, 8, 16, 16], start=1):
        items.column_dimensions[get_column_letter(i)].width = width

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def to_docx(records: list[dict[str, Any]], title: str) -> bytes:
    """Word document: heading, summary line, and one table of documents."""
    from docx import Document as DocxDocument
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    rows = [_report_row(r) for r in records]
    doc = DocxDocument()
    doc.add_heading(title, level=1)
    doc.add_paragraph(
        f"{len(rows)} approved document{'' if len(rows) == 1 else 's'} · "
        f"generated {date.today().isoformat()}"
    )

    table = doc.add_table(rows=1, cols=len(_REPORT_COLUMNS))
    table.style = "Light Grid Accent 1"
    for cell, (header, _, _) in zip(table.rows[0].cells, _REPORT_COLUMNS):
        cell.text = header
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(9)

    for row in rows:
        cells = table.add_row().cells
        for cell, (_, key, is_money) in zip(cells, _REPORT_COLUMNS):
            cell.text = _fmt_amount(row[key]) if is_money else str(row[key])
            paragraph = cell.paragraphs[0]
            if is_money:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in paragraph.runs:
                run.font.size = Pt(9)

    if rows:
        totals = _money_totals(rows)
        cells = table.add_row().cells
        for i, (cell, (_, key, is_money)) in enumerate(zip(cells, _REPORT_COLUMNS)):
            cell.text = "TOTAL" if i == 0 else (_fmt_amount(totals[key]) if is_money else "")
            paragraph = cell.paragraphs[0]
            if is_money:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def to_pdf(records: list[dict[str, Any]], title: str) -> bytes:
    """PDF report, laid out with PyMuPDF's Story engine.

    Story takes HTML/CSS and paginates it for us, so a 200-row export flows
    onto as many A4 pages as it needs without any manual y-cursor arithmetic.
    PyMuPDF is already a dependency (PDF rasterization), so no extra library.
    """
    import fitz

    rows = [_report_row(r) for r in records]

    def esc(value: Any) -> str:
        return (
            str(value)
            .replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
        )

    body = ["<tr>" + "".join(f"<th>{esc(h)}</th>" for h, _, _ in _REPORT_COLUMNS) + "</tr>"]
    for row in rows:
        body.append(
            "<tr>"
            + "".join(
                f'<td class="{"num" if is_money else "txt"}">'
                f"{esc(_fmt_amount(row[key]) if is_money else row[key])}</td>"
                for _, key, is_money in _REPORT_COLUMNS
            )
            + "</tr>"
        )
    if rows:
        totals = _money_totals(rows)
        body.append(
            '<tr class="total">'
            + "".join(
                f'<td class="{"num" if is_money else "txt"}">'
                f"{esc(_fmt_amount(totals[key])) if is_money else ('TOTAL' if i == 0 else '')}</td>"
                for i, (_, key, is_money) in enumerate(_REPORT_COLUMNS)
            )
            + "</tr>"
        )

    html = f"""
      <html><body>
        <h2>{esc(title)}</h2>
        <p class="sub">{len(rows)} approved document{'' if len(rows) == 1 else 's'}
           &#183; generated {date.today().isoformat()}</p>
        <table>{''.join(body)}</table>
      </body></html>
    """
    css = """
      body { font-family: sans-serif; font-size: 8pt; color: #1b1b1f; }
      h2 { font-size: 14pt; margin-bottom: 2pt; }
      p.sub { color: #55565c; font-size: 8pt; margin-top: 0; }
      th { background-color: #e8eaf6; text-align: left; font-weight: bold; padding: 3pt; }
      td { padding: 3pt; }
      td.num { text-align: right; }
      tr.total td { font-weight: bold; }
    """

    story = fitz.Story(html=html, user_css=css)
    buf = io.BytesIO()
    writer = fitz.DocumentWriter(buf)
    page = fitz.paper_rect("a4")
    frame = page + (36, 36, -36, -36)  # ~12.7mm margins
    more = True
    while more:
        device = writer.begin_page(page)
        more, _ = story.place(frame)
        story.draw(device)
        writer.end_page()
    writer.close()
    return buf.getvalue()


# --- format registry ---------------------------------------------------------

# fmt -> (builder(records, title) -> bytes, media type). The route validates the
# requested format against these keys, so adding a format here is all it takes
# to expose it in the API (and the frontend's picker mirrors this list).
REPORT_FORMATS: dict[str, tuple[Callable[[list[dict[str, Any]], str], bytes], str]] = {
    "pdf": (to_pdf, "application/pdf"),
    "docx": (
        to_docx,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ),
    "xlsx": (
        to_xlsx,
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    ),
    "csv": (lambda recs, _title: to_csv_many(recs).encode("utf-8-sig"), "text/csv"),
    "json": (lambda recs, _title: to_json_many(recs).encode("utf-8"), "application/json"),
}
